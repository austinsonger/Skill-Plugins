"""
Generic Compliance Evidence Package Builder

Reads a single JSON file describing (framework, CSP, controls, ...) and
produces a matching DOCX narrative + XLSX evidence matrix.

Usage:
    python3 build_evidence_package.py path/to/package_data.json [--out-dir DIR]

Schema (see references/data_schema.md for the full spec):

{
  "framework": {
    "name":        "HIPAA Security & Breach Notification Rules",
    "short_name":  "HIPAA",
    "scope":       "45 CFR Parts 160 & 164 ...",
    "id_label":    "Citation",          # column header for control ID
    "type_column": {                    # OPTIONAL; only if framework has R/A or similar
        "header":      "Type (R/A)",
        "values": {                     # color hints for the XLSX type cell
            "R": {"label": "Required",    "fill": "C00000", "text": "FFFFFF"},
            "A": {"label": "Addressable", "fill": "FFC000", "text": "000000"}
        }
    },
    "authoritative_source": "45 CFR ...",
    "framework_notes": [ "...", "..." ]  # bullets that appear in the Evidence Sourcing section
  },
  "csp": {
    "name":           "Amazon Web Services (AWS)",
    "short_name":     "AWS",
    "shared_resp":    "AWS is responsible for ...; customer is responsible for ...",
    "attestation_repo": {                # where inherited evidence comes from
        "name": "AWS Artifact",
        "url":  "https://aws.amazon.com/artifact"
    },
    "baa_or_agreement": {                # OPTIONAL; only if framework needs one
        "name": "AWS Business Associate Addendum",
        "where_signed": "AWS Artifact -> Agreements -> AWS BAA"
    },
    "foundational_services": [           # ~8-12 services that recur across controls
        {"name": "AWS CloudTrail", "blurb": "API audit log. ..."},
        ...
    ],
    "console_term":   "Console Path / CLI"  # column header for portal/CLI path
  },
  "deliverable_name_stem": "HIPAA_AWS",     # used for the output filenames
  "categories": [
    {
      "name":  "Administrative Safeguards (§164.308)",
      "rows": [
        {
          "id":           "164.308(a)(1)(i)",      # control / criterion ID
          "name":         "Security Management Process",
          "type":         "R",                     # OPTIONAL; only if type_column set
          "description":  "Implement policies ...",
          "services":     "AWS Security Hub; ...",
          "evidence":     "Information security ...",
          "path":         "Security Hub -> Settings",
          "frequency":    "Annual + continuous",
          "owner":        "Security / Compliance"
        },
        ...
      ]
    },
    ...
  ],
  "inherited": [                                # OPTIONAL but recommended
    {"id": "164.310(a)(1)", "why": "...", "evidence": "..."},
    ...
  ],
  "purpose_paragraph": "This guide maps ...",
  "evidence_sourcing_principles": [             # bullets in the Evidence Sourcing section
    {"name": "Six-year retention.", "body": "..."}
  ],
  "evidence_retention_bullets": [               # bullets in the Retention/Handoff section
    "Retain ...",
    ...
  ],
  "references": [                               # bullets in the References section
    "AICPA ...",
    ...
  ],
  "extra_xlsx_tabs": [                          # OPTIONAL extra sheets
    {"title": "Required vs Addressable", "rows": [["Required (R)", "..."], ...]}
  ]
}

Output:
  <out-dir>/<stem>_Evidence_Guide.docx
  <out-dir>/<stem>_Evidence_Matrix.xlsx
"""

from __future__ import annotations

import argparse
import json
import os
import subprocess
import sys

# -------- Install deps quietly --------
for pkg in ("python-docx", "openpyxl"):
    try:
        __import__(pkg.replace("-", "_"))
    except ImportError:
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", pkg,
             "--break-system-packages", "--quiet"],
        )

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.formatting.rule import CellIsRule

# -------- Palette (shared across all CSPs/frameworks) --------
PALETTE = {
    "AWS":   {"primary": "142E57", "accent": "C25710", "soft_bg": "FFF4E5"},
    "Azure": {"primary": "003B6A", "accent": "0078D4", "soft_bg": "E5F1FB"},
    "GCP":   {"primary": "1A47A3", "accent": "4285F4", "soft_bg": "E8F0FE"},
    "OCI":   {"primary": "8B0000", "accent": "C74634", "soft_bg": "FBE9E7"},
    "IBM":   {"primary": "002D9C", "accent": "0F62FE", "soft_bg": "EDF5FF"},
    "DEFAULT": {"primary": "142E57", "accent": "555555", "soft_bg": "F0F0F0"},
}

AMBER_BG = "FFF4E5"
AMBER_HDR = "B26A00"

STATUS_VALUES = ["Not Started", "In Progress", "Evidence Collected", "Reviewed", "N/A"]
RISK_VALUES = ["Critical", "High", "Medium", "Low", "Informational"]
STATUS_FILLS = {
    "Not Started": "F2F2F2",
    "In Progress": "FFC000",
    "Evidence Collected": "BDD7EE",
    "Reviewed": "70AD47",
    "N/A": "BFBFBF",
}
RISK_FILLS = {
    "Critical": ("C00000", "FFFFFF", True),
    "High":     ("ED7D31", "FFFFFF", True),
    "Medium":   ("FFC000", "000000", True),
    "Low":      ("70AD47", "FFFFFF", True),
    "Informational": ("BFBFBF", "000000", False),
}


# =========================================================================
# Helpers
# =========================================================================

def _palette(csp_short):
    return PALETTE.get(csp_short.upper(), PALETTE["DEFAULT"])

def _add_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _set_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "B0B0B0")
        borders.append(b)
    tcPr.append(borders)

def _hex_to_rgb(hex_color):
    return RGBColor(int(hex_color[0:2], 16),
                    int(hex_color[2:4], 16),
                    int(hex_color[4:6], 16))


# =========================================================================
# DOCX builder
# =========================================================================

def build_docx(data, out_path):
    fw = data["framework"]
    csp = data["csp"]
    pal = _palette(csp["short_name"])
    PRIMARY = _hex_to_rgb(pal["primary"])
    ACCENT = _hex_to_rgb(pal["accent"])

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.left_margin = Inches(0.8); sec.right_margin = Inches(0.8)
    sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.8)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # Cover ----------------------------------------------------------------
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(fw["name"])
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = PRIMARY

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{csp['name']} Evidence Collection Guide")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = ACCENT

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{fw.get('scope','')}\nVersion {data.get('version','1.0')}  |  {data.get('date','')}")
    r.font.size = Pt(11); r.italic = True

    doc.add_paragraph(); doc.add_paragraph()

    # Purpose callout
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False; tbl.columns[0].width = Inches(6.9)
    cell = tbl.rows[0].cells[0]
    _add_shading(cell, pal["soft_bg"])
    p = cell.paragraphs[0]
    r = p.add_run("Purpose"); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = ACCENT
    cell.add_paragraph(data["purpose_paragraph"])

    doc.add_page_break()

    # Section 1: Shared Responsibility
    h = doc.add_heading(f"1. {csp['short_name']} Shared Responsibility Model", level=1)
    for run in h.runs: run.font.color.rgb = PRIMARY
    doc.add_paragraph(csp["shared_resp"])

    if data.get("inherited"):
        doc.add_paragraph(
            "Specific citations evidenced via inherited controls (see the matrix's "
            f"'Inherited from {csp['short_name']}' tab):"
        )
        for inh in data["inherited"]:
            doc.add_paragraph(f"{inh['id']} — {inh['why']}", style="List Bullet")

    doc.add_paragraph()

    # Section 2: Evidence sourcing principles
    h = doc.add_heading("2. Evidence Sourcing Principles", level=1)
    for run in h.runs: run.font.color.rgb = PRIMARY
    for principle in data.get("evidence_sourcing_principles", []):
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(principle["name"]); r.bold = True
        p.add_run(" " + principle["body"])

    doc.add_paragraph()

    # Section 3: Foundational services
    h = doc.add_heading(f"3. Foundational {csp['short_name']} Services", level=1)
    for run in h.runs: run.font.color.rgb = PRIMARY
    for svc in csp.get("foundational_services", []):
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(svc["name"]); r.bold = True
        p.add_run(" — " + svc["blurb"])

    doc.add_page_break()

    # Section 4: Per-control detail ------------------------------------------
    h = doc.add_heading("4. Control-by-Control Evidence Detail", level=1)
    for run in h.runs: run.font.color.rgb = PRIMARY
    doc.add_paragraph(
        f"Each entry below names the {fw['id_label']}, the standard or specification, "
        + ("the type flag, " if fw.get("type_column") else "")
        + "a plain-language description, the services that produce evidence, the artifact to "
        + "retain, the console/CLI path, and the cadence. Use the companion XLSX workbook to "
        + "track status across the audit period."
    )

    for cat in data["categories"]:
        doc.add_paragraph()
        h2 = doc.add_heading(cat["name"], level=2)
        for run in h2.runs: run.font.color.rgb = PRIMARY

        for row in cat["rows"]:
            tbl = doc.add_table(rows=1, cols=2)
            tbl.autofit = False
            tbl.columns[0].width = Inches(1.5)
            tbl.columns[1].width = Inches(5.4)
            c_id = tbl.rows[0].cells[0]; c_text = tbl.rows[0].cells[1]
            _add_shading(c_id, pal["primary"])
            _add_shading(c_text, pal["soft_bg"])
            _set_borders(c_id); _set_borders(c_text)

            label = row["id"]
            if fw.get("type_column") and row.get("type"):
                label = f"{label}\n[{row['type']}]"
            p = c_id.paragraphs[0]
            r = p.add_run(label); r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.size = Pt(11)

            p = c_text.paragraphs[0]
            r = p.add_run(row["name"]); r.bold = True; r.font.size = Pt(11)
            c_text.add_paragraph().add_run(row["description"]).italic = True

            for field, label in [
                ("services", f"{csp['short_name']} services"),
                ("evidence", "Evidence to retain"),
                ("path", "Where to find it"),
            ]:
                if row.get(field):
                    p = doc.add_paragraph()
                    r = p.add_run(f"{label}: "); r.bold = True
                    p.add_run(row[field])

            p = doc.add_paragraph()
            if row.get("frequency"):
                r = p.add_run("Cadence: "); r.bold = True
                p.add_run(row["frequency"])
            if row.get("owner"):
                r = p.add_run("    |    Suggested owner: "); r.bold = True
                p.add_run(row["owner"])

            doc.add_paragraph()

    doc.add_page_break()

    # Section 5: Retention & handoff
    h = doc.add_heading("5. Evidence Retention and Auditor Handoff", level=1)
    for run in h.runs: run.font.color.rgb = PRIMARY
    for item in data.get("evidence_retention_bullets", []):
        doc.add_paragraph(item, style="List Bullet")

    # Section 6: References
    h = doc.add_heading("6. References", level=1)
    for run in h.runs: run.font.color.rgb = PRIMARY
    for item in data.get("references", []):
        doc.add_paragraph(item, style="List Bullet")

    doc.save(out_path)
    print(f"Wrote {out_path}")


# =========================================================================
# XLSX builder
# =========================================================================

def build_xlsx(data, out_path):
    fw = data["framework"]
    csp = data["csp"]
    pal = _palette(csp["short_name"])
    PRIMARY = pal["primary"]
    ACCENT = pal["accent"]

    wb = Workbook()

    # ----- Cover -----
    cov = wb.active; cov.title = "Cover"
    cov.column_dimensions["A"].width = 28
    cov.column_dimensions["B"].width = 80

    cov["A1"] = f"{fw['short_name']} — {csp['short_name']} Evidence Matrix"
    cov["A1"].font = Font(name="Calibri", size=20, bold=True, color=PRIMARY)
    cov.merge_cells("A1:B1")

    meta_rows = [
        ("Framework", fw.get("name", "")),
        ("Scope", fw.get("scope", "")),
        ("Environment", csp.get("name", "")),
        ("Version", data.get("version", "1.0")),
        ("Date", data.get("date", "")),
        ("Sheets", "Cover | Evidence Matrix | Inherited from " + csp["short_name"]
                   + (" | " + " | ".join(t["title"] for t in data.get("extra_xlsx_tabs", []))
                      if data.get("extra_xlsx_tabs") else "")
                   + " | Legend & Summary"),
    ]
    for i, (k, v) in enumerate(meta_rows, 3):
        cov.cell(row=i, column=1, value=k).font = Font(bold=True, color=PRIMARY)
        cov.cell(row=i, column=2, value=v).alignment = Alignment(wrap_text=True, vertical="top")

    # Instructions
    instr_start = 3 + len(meta_rows) + 1
    cov.cell(row=instr_start, column=1, value="How to use this workbook").font = \
        Font(bold=True, size=12, color=ACCENT)
    cov.merge_cells(start_row=instr_start, start_column=1, end_row=instr_start, end_column=2)
    for i, line in enumerate(data.get("xlsx_instructions", [
        "1. Assign an owner to each row in the Evidence Matrix sheet.",
        "2. Configure the named service and validate it is in scope under the relevant agreement / BAA.",
        "3. Capture the artifact on the cadence shown and store it in your evidence repository.",
        "4. Update Status (dropdown) and paste the storage path or URL into Evidence Location.",
        "5. The Legend & Summary sheet auto-counts status. Review it weekly during fieldwork.",
    ])):
        cov.cell(row=instr_start + 1 + i, column=1, value=line)
        cov.merge_cells(start_row=instr_start + 1 + i, start_column=1,
                        end_row=instr_start + 1 + i, end_column=2)

    # Amber derivation note
    note_row = instr_start + 1 + len(data.get("xlsx_instructions", [1]*5)) + 1
    cov.cell(row=note_row, column=1, value="Evidence Derivation Note").font = \
        Font(bold=True, color=AMBER_HDR)
    cov.cell(row=note_row, column=1).fill = PatternFill(start_color=AMBER_BG, fill_type="solid")
    cov.merge_cells(start_row=note_row, start_column=1, end_row=note_row, end_column=2)
    cov.cell(row=note_row + 1, column=1,
             value=data.get("derivation_note",
                            f"Evidence statements are anchored to {fw.get('authoritative_source','the framework')} "
                            f"(framework-native source) and {csp['short_name']}-published implementation guidance. "
                            f"Validate console paths and service availability before issuing to an auditor."))
    cov.cell(row=note_row + 1, column=1).fill = PatternFill(start_color=AMBER_BG, fill_type="solid")
    cov.cell(row=note_row + 1, column=1).alignment = Alignment(wrap_text=True, vertical="top")
    cov.merge_cells(start_row=note_row + 1, start_column=1, end_row=note_row + 1, end_column=2)
    cov.row_dimensions[note_row + 1].height = 90

    # ----- Evidence Matrix -----
    em = wb.create_sheet("Evidence Matrix")
    type_col = fw.get("type_column")

    headers = ["Category", fw.get("id_label", "ID"), "Standard / Specification"]
    if type_col: headers.append(type_col["header"])
    headers += [
        "Description",
        f"{csp['short_name']} Services (Evidence Source)",
        "Evidence to Retain",
        csp.get("console_term", "Console Path / CLI"),
        "Cadence", "Owner",
        "Status", "Risk Level", "Evidence Location (path/URL)", "Notes",
    ]

    # Header row
    for col, name in enumerate(headers, 1):
        c = em.cell(row=1, column=col, value=name)
        c.font = Font(bold=True, color="FFFFFF", size=11)
        c.fill = PatternFill(start_color=PRIMARY, fill_type="solid")
        c.alignment = Alignment(wrap_text=True, vertical="center", horizontal="center")

    # Find amber-highlight columns (services + evidence)
    services_col = headers.index(f"{csp['short_name']} Services (Evidence Source)") + 1
    evidence_col = headers.index("Evidence to Retain") + 1
    em.cell(row=1, column=services_col).fill = PatternFill(start_color=AMBER_HDR, fill_type="solid")
    em.cell(row=1, column=evidence_col).fill = PatternFill(start_color=AMBER_HDR, fill_type="solid")

    # Column widths
    base_widths = [24, 18, 32]
    if type_col: base_widths.append(10)
    base_widths += [44, 36, 56, 36, 22, 22, 14, 14, 36, 28]
    for i, w in enumerate(base_widths, 1):
        em.column_dimensions[get_column_letter(i)].width = w

    # Data rows
    row_num = 2
    for cat in data["categories"]:
        for ctrl in cat["rows"]:
            col = 1
            em.cell(row=row_num, column=col, value=cat["name"]); col += 1
            em.cell(row=row_num, column=col, value=ctrl["id"]); col += 1
            em.cell(row=row_num, column=col, value=ctrl["name"]); col += 1
            if type_col:
                tval = ctrl.get("type", "")
                tcell = em.cell(row=row_num, column=col, value=tval)
                tdef = type_col["values"].get(tval, {})
                if tdef:
                    tcell.fill = PatternFill(start_color=tdef.get("fill", "FFFFFF"), fill_type="solid")
                    tcell.font = Font(color=tdef.get("text", "000000"), bold=True)
                    tcell.alignment = Alignment(horizontal="center", vertical="center")
                col += 1
            em.cell(row=row_num, column=col, value=ctrl.get("description", "")); col += 1
            em.cell(row=row_num, column=col, value=ctrl.get("services", "")); col += 1
            em.cell(row=row_num, column=col, value=ctrl.get("evidence", "")); col += 1
            em.cell(row=row_num, column=col, value=ctrl.get("path", "")); col += 1
            em.cell(row=row_num, column=col, value=ctrl.get("frequency", "")); col += 1
            em.cell(row=row_num, column=col, value=ctrl.get("owner", "")); col += 1
            em.cell(row=row_num, column=col, value="Not Started"); col += 1
            em.cell(row=row_num, column=col, value=ctrl.get("default_risk", "Medium")); col += 1
            em.cell(row=row_num, column=col, value=""); col += 1
            em.cell(row=row_num, column=col, value="")

            # Amber fill on services + evidence
            em.cell(row=row_num, column=services_col).fill = \
                PatternFill(start_color=AMBER_BG, fill_type="solid")
            em.cell(row=row_num, column=evidence_col).fill = \
                PatternFill(start_color=AMBER_BG, fill_type="solid")

            for c in range(1, len(headers) + 1):
                em.cell(row=row_num, column=c).alignment = Alignment(wrap_text=True, vertical="top")
            if type_col:
                type_col_idx = headers.index(type_col["header"]) + 1
                em.cell(row=row_num, column=type_col_idx).alignment = \
                    Alignment(horizontal="center", vertical="center")

            row_num += 1

    last_row = row_num - 1
    status_col_idx = headers.index("Status") + 1
    risk_col_idx = headers.index("Risk Level") + 1
    status_letter = get_column_letter(status_col_idx)
    risk_letter = get_column_letter(risk_col_idx)

    # Validations
    status_dv = DataValidation(
        type="list", formula1='"' + ",".join(STATUS_VALUES) + '"', allow_blank=True)
    status_dv.add(f"{status_letter}2:{status_letter}{last_row}")
    em.add_data_validation(status_dv)

    risk_dv = DataValidation(
        type="list", formula1='"' + ",".join(RISK_VALUES) + '"', allow_blank=True)
    risk_dv.add(f"{risk_letter}2:{risk_letter}{last_row}")
    em.add_data_validation(risk_dv)

    # Conditional formatting
    for status, fill in STATUS_FILLS.items():
        em.conditional_formatting.add(
            f"{status_letter}2:{status_letter}{last_row}",
            CellIsRule(operator="equal", formula=[f'"{status}"'],
                       fill=PatternFill(start_color=fill, fill_type="solid"),
                       font=Font(color="FFFFFF", bold=True) if status == "Reviewed" else None))

    for risk, (fill, text, bold) in RISK_FILLS.items():
        em.conditional_formatting.add(
            f"{risk_letter}2:{risk_letter}{last_row}",
            CellIsRule(operator="equal", formula=[f'"{risk}"'],
                       fill=PatternFill(start_color=fill, fill_type="solid"),
                       font=Font(color=text, bold=bold)))

    em.freeze_panes = "C2"

    # ----- Inherited tab -----
    if data.get("inherited"):
        inh = wb.create_sheet(f"Inherited from {csp['short_name']}")
        inh.column_dimensions["A"].width = 18
        inh.column_dimensions["B"].width = 55
        inh.column_dimensions["C"].width = 60
        inh["A1"] = fw.get("id_label", "Citation")
        inh["B1"] = "Why Inherited"
        inh["C1"] = f"Evidence ({csp.get('attestation_repo',{}).get('name','attestation portal')})"
        for col in (1, 2, 3):
            c = inh.cell(row=1, column=col)
            c.font = Font(bold=True, color="FFFFFF")
            c.fill = PatternFill(start_color=PRIMARY, fill_type="solid")
        for i, row in enumerate(data["inherited"], 2):
            inh.cell(row=i, column=1, value=row["id"]).font = Font(bold=True)
            inh.cell(row=i, column=2, value=row["why"]).alignment = \
                Alignment(wrap_text=True, vertical="top")
            inh.cell(row=i, column=3, value=row["evidence"]).alignment = \
                Alignment(wrap_text=True, vertical="top")
            inh.row_dimensions[i].height = 55

    # ----- Extra tabs (framework-specific) -----
    for extra in data.get("extra_xlsx_tabs", []):
        sh = wb.create_sheet(extra["title"][:31])
        sh.column_dimensions["A"].width = 24
        sh.column_dimensions["B"].width = 88
        if extra.get("headers"):
            for col, name in enumerate(extra["headers"], 1):
                c = sh.cell(row=1, column=col, value=name)
                c.font = Font(bold=True, color="FFFFFF")
                c.fill = PatternFill(start_color=PRIMARY, fill_type="solid")
        for i, row in enumerate(extra["rows"], 2 if extra.get("headers") else 1):
            for col, val in enumerate(row, 1):
                cell = sh.cell(row=i, column=col, value=val)
                cell.alignment = Alignment(wrap_text=True, vertical="top")
                if col == 1: cell.font = Font(bold=True)
            sh.row_dimensions[i].height = 50

    # ----- Legend & Summary -----
    leg = wb.create_sheet("Legend & Summary")
    leg.column_dimensions["A"].width = 24
    leg.column_dimensions["B"].width = 60
    leg.column_dimensions["D"].width = 24
    leg.column_dimensions["E"].width = 16

    leg["A1"] = "Status Values"
    leg["A1"].font = Font(bold=True, size=12, color=PRIMARY)
    for i, (s, desc) in enumerate([
        ("Not Started", "Owner assigned; evidence not yet collected."),
        ("In Progress", "Owner is collecting or service is being configured."),
        ("Evidence Collected", "Artifact stored in evidence repository, awaiting QA review."),
        ("Reviewed", "QA-reviewed and confirmed; ready for auditor."),
        ("N/A", "Not applicable to this environment (justify in Notes)."),
    ], 2):
        leg.cell(row=i, column=1, value=s).font = Font(bold=True)
        leg.cell(row=i, column=2, value=desc)

    leg["A9"] = "Risk Levels"
    leg["A9"].font = Font(bold=True, size=12, color=PRIMARY)
    for i, (s, desc) in enumerate([
        ("Critical", "Failure would directly cause an audit qualification or material breach."),
        ("High", "Failure likely to be reported as a finding."),
        ("Medium", "Likely deviation; remediable within the audit period."),
        ("Low", "Minor deviation; documentation-level fix."),
        ("Informational", "Observation only; not a finding."),
    ], 10):
        leg.cell(row=i, column=1, value=s).font = Font(bold=True)
        leg.cell(row=i, column=2, value=desc)

    # Summary counts
    leg["D1"] = "Status Summary"
    leg["D1"].font = Font(bold=True, size=12, color=PRIMARY)
    for i, s in enumerate(STATUS_VALUES, 2):
        leg.cell(row=i, column=4, value=s).font = Font(bold=True)
        leg.cell(row=i, column=5,
                 value=f'=COUNTIF(\'Evidence Matrix\'!{status_letter}2:{status_letter}{last_row},"{s}")')

    leg["D9"] = "Risk Summary"
    leg["D9"].font = Font(bold=True, size=12, color=PRIMARY)
    for i, s in enumerate(RISK_VALUES, 10):
        leg.cell(row=i, column=4, value=s).font = Font(bold=True)
        leg.cell(row=i, column=5,
                 value=f'=COUNTIF(\'Evidence Matrix\'!{risk_letter}2:{risk_letter}{last_row},"{s}")')

    # Type summary (if framework has a type column)
    if type_col:
        type_col_letter = get_column_letter(headers.index(type_col["header"]) + 1)
        leg["D17"] = f"{type_col['header']} Summary"
        leg["D17"].font = Font(bold=True, size=12, color=PRIMARY)
        row = 18
        for tval, tdef in type_col["values"].items():
            cell = leg.cell(row=row, column=4, value=tdef.get("label", tval))
            cell.font = Font(bold=True, color=tdef.get("fill", "000000"))
            leg.cell(row=row, column=5,
                     value=f'=COUNTIF(\'Evidence Matrix\'!{type_col_letter}2:{type_col_letter}{last_row},"{tval}")')
            row += 1
        leg.cell(row=row, column=4, value="Total Controls").font = \
            Font(bold=True, color=PRIMARY)
        leg.cell(row=row, column=5, value=last_row - 1)
    else:
        leg["D17"] = "Total Controls"
        leg["D17"].font = Font(bold=True, color=PRIMARY)
        leg["E17"] = last_row - 1

    wb.save(out_path)
    print(f"Wrote {out_path}")


# =========================================================================
# Entry point
# =========================================================================

def main():
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    ap.add_argument("data_json", help="Path to the package data JSON file")
    ap.add_argument("--out-dir", default=".",
                    help="Output directory (default: current dir)")
    args = ap.parse_args()

    with open(args.data_json, "r") as f:
        data = json.load(f)

    stem = data.get("deliverable_name_stem")
    if not stem:
        fw_short = data["framework"]["short_name"].replace(" ", "_")
        csp_short = data["csp"]["short_name"].replace(" ", "_")
        stem = f"{fw_short}_{csp_short}"

    os.makedirs(args.out_dir, exist_ok=True)
    docx_path = os.path.join(args.out_dir, f"{stem}_Evidence_Guide.docx")
    xlsx_path = os.path.join(args.out_dir, f"{stem}_Evidence_Matrix.xlsx")

    build_docx(data, docx_path)
    build_xlsx(data, xlsx_path)
    print("Done.")


if __name__ == "__main__":
    main()
