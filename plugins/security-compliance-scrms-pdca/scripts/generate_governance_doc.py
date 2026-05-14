#!/usr/bin/env python3
"""
SCRMS-PDCA: Governance Document Generator
Generates a Word document with a Policy + Standard + Guideline shell for a given control domain.
Follows the SCRMS Principle #4 governance documentation hierarchy.

Usage:
    python generate_governance_doc.py --domain <domain> --org-name <name> --output <path>
                                      [--regulatory-drivers <req1,req2>]

Example:
    python generate_governance_doc.py \\
        --domain "Access Management" \\
        --org-name "Acme Corp" \\
        --regulatory-drivers "HIPAA,SOC 2 Type II,ISO 27001" \\
        --output access_management_policy.docx
"""

import argparse
import sys
from datetime import date

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx",
                    "--break-system-packages", "-q"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement


DOMAIN_TEMPLATES = {
    "Access Management": {
        "policy_statement": (
            "All access to {org_name} information systems, applications, and data must be authorized, "
            "granted on the principle of least privilege, and revoked promptly upon termination or role change. "
            "Privileged access must be subject to enhanced controls including multi-factor authentication and periodic review."
        ),
        "scope": "All employees, contractors, vendors, and third parties who access {org_name} systems or data.",
        "policy_requirements": [
            "All user accounts must be formally requested, approved, and provisioned through the identity management process.",
            "Access rights must reflect the minimum permissions required to perform assigned duties (least privilege).",
            "Privileged accounts (administrators, root, service accounts) must be managed through a Privileged Access Management (PAM) process.",
            "Multi-factor authentication (MFA) is required for all remote access and all privileged account use.",
            "User access must be reviewed on a periodic basis to validate continued business need.",
            "Access must be revoked within 24 hours of employment termination or role change.",
            "Shared/generic accounts are prohibited except where technically required, with compensating controls documented.",
        ],
        "standard_controls": [
            ("Identity Provisioning", "New accounts must be provisioned via a formal request approved by the user's manager and system owner. Requests must be logged in the identity management system."),
            ("Authentication Requirements", "Passwords must meet complexity requirements (minimum 12 characters, mix of character types). MFA is required for all remote access sessions and all privileged account usage."),
            ("Privileged Access Management", "All privileged accounts must be enrolled in the PAM solution. Privileged session recording is required for critical systems. Passwords for shared privileged accounts must be rotated after each use."),
            ("Access Reviews", "Access reviews must be conducted quarterly for privileged accounts and annually for standard accounts. Results must be documented and exceptions remediated within 30 days."),
            ("Termination Procedures", "HR must notify IT within 2 hours of an employee's last day. All accounts must be disabled within 24 hours and fully removed within 30 days."),
        ],
        "guidelines": [
            "Consider implementing Just-In-Time (JIT) access for privileged systems to minimize standing access exposure.",
            "Role-Based Access Control (RBAC) is the recommended model for new application deployments.",
            "For high-security environments, consider hardware security keys (FIDO2) as the MFA method.",
            "Service accounts should use Managed Identity or workload identity federation where supported, rather than stored credentials.",
        ],
        "control_objectives": [
            "CO-AC-1: Ensure all user accounts are authorized, documented, and managed throughout their lifecycle.",
            "CO-AC-2: Ensure access rights reflect least-privilege principles and are periodically validated.",
            "CO-AC-3: Ensure privileged access is controlled, monitored, and time-limited.",
            "CO-AC-4: Ensure authentication mechanisms meet minimum assurance levels appropriate to asset sensitivity.",
        ],
    },
    "Data Protection": {
        "policy_statement": (
            "{org_name} is committed to protecting sensitive data — including personal information, "
            "protected health information, financial data, and intellectual property — throughout its "
            "lifecycle. Data must be classified, handled, stored, transmitted, and disposed of in "
            "accordance with its sensitivity and applicable regulatory requirements."
        ),
        "scope": "All data created, collected, stored, processed, or transmitted by {org_name} or on its behalf.",
        "policy_requirements": [
            "All data must be classified according to the {org_name} Data Classification Policy.",
            "Sensitive data must be encrypted at rest using approved cryptographic algorithms.",
            "Sensitive data must be encrypted in transit using approved protocols.",
            "Data retention schedules must be documented and enforced; data must be securely disposed of upon reaching end-of-retention.",
            "Personal data processing must be documented in the Records of Processing Activities (ROPA).",
            "Privacy Impact Assessments (PIA/DPIA) must be conducted for new processing activities involving personal data.",
            "Data Loss Prevention (DLP) controls must be implemented to prevent unauthorized exfiltration of sensitive data.",
        ],
        "standard_controls": [
            ("Data Classification", "Data must be classified into one of four tiers: Public, Internal, Confidential, or Restricted. Classification labels must be applied to documents and storage locations."),
            ("Encryption at Rest", "Confidential and Restricted data must be encrypted at rest using AES-256. Encryption keys must be managed separately from encrypted data."),
            ("Encryption in Transit", "All transmission of Confidential or Restricted data must use TLS 1.2 or higher. Unencrypted protocols (HTTP, FTP, Telnet) are prohibited for sensitive data."),
            ("Data Retention", "Retention periods must be defined per data type in the Data Retention Schedule. Data must be securely destroyed using NIST SP 800-88 guidelines when retention expires."),
            ("Privacy Compliance", "A ROPA must be maintained and reviewed annually. DPIAs must be completed before initiating new personal data processing activities."),
        ],
        "guidelines": [
            "Consider data minimization principles — collect only what is necessary for the stated purpose.",
            "Tokenization or pseudonymization is recommended for analytics workloads involving personal data.",
            "Cloud storage buckets and object stores should be configured with explicit deny-public access policies by default.",
            "Data classification tools (auto-labeling) can reduce classification burden on end users in large environments.",
        ],
        "control_objectives": [
            "CO-DP-1: Ensure sensitive data is identifiable, classified, and labeled appropriately.",
            "CO-DP-2: Ensure sensitive data is protected from unauthorized access at rest and in transit.",
            "CO-DP-3: Ensure personal data processing is lawful, documented, and compliant with applicable privacy regulations.",
            "CO-DP-4: Ensure data is retained only as long as necessary and disposed of securely.",
        ],
    },
}

# Generic template for domains not in the library
GENERIC_TEMPLATE = {
    "policy_statement": (
        "{org_name} requires that {domain} controls are implemented, maintained, and operating effectively "
        "to protect organizational assets and meet applicable compliance obligations. All personnel with "
        "responsibilities in this domain must comply with this policy."
    ),
    "scope": "All {org_name} personnel, systems, and processes within the {domain} control domain.",
    "policy_requirements": [
        f"[Requirement 1 — describe a mandatory control or behavior for this domain]",
        f"[Requirement 2 — describe a mandatory control or behavior for this domain]",
        f"[Requirement 3 — describe a mandatory control or behavior for this domain]",
        f"[Requirement 4 — describe a mandatory control or behavior for this domain]",
        f"[Requirement 5 — describe exception handling and approval process]",
    ],
    "standard_controls": [
        ("[Control Area 1]", "[Specific, measurable requirement. Describe exactly what must be done, by whom, and how frequently.]"),
        ("[Control Area 2]", "[Specific, measurable requirement. Include any thresholds, timelines, or technical specifications.]"),
        ("[Control Area 3]", "[Specific, measurable requirement. Reference any tools, procedures, or evidence requirements.]"),
    ],
    "guidelines": [
        "[Best practice recommendation — not mandatory, but recommended for this domain]",
        "[Technology or process option to consider when implementing these controls]",
        "[Maturity advancement idea — how to move from CMM 3 to CMM 4 in this area]",
    ],
    "control_objectives": [
        f"CO-1: [Outcome that must be achieved — what does 'implemented' look like for this domain?]",
        f"CO-2: [Second outcome — typically corresponds to a different aspect of the domain]",
        f"CO-3: [Third outcome — consider PPTDF scope: People, Process, Technology, Data, Facilities]",
    ],
}


def set_heading_style(para, text, level, color=None):
    para.clear()
    run = para.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_metadata_table(doc, org_name, domain, version, reg_drivers):
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    fields = [
        ("Organization:", org_name),
        ("Policy Domain:", domain),
        ("Version / Date:", f"{version} — {date.today().strftime('%B %d, %Y')}"),
        ("Regulatory Drivers:", reg_drivers or "To be determined"),
        ("SCRMS Phase / Principle:", "PLAN — Principle #4: Publish Governance Documentation"),
    ]
    for i, (label, value) in enumerate(fields):
        table.cell(i, 0).text = label
        table.cell(i, 0).paragraphs[0].runs[0].bold = True
        table.cell(i, 1).text = value
    doc.add_paragraph()


def add_section(doc, title, level=2):
    p = doc.add_paragraph()
    set_heading_style(p, title, level)


def add_bulleted_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_two_col_table(doc, rows_data, col1_header, col2_header):
    table = doc.add_table(rows=1 + len(rows_data), cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = col1_header
    hdr[1].text = col2_header
    for cell in hdr:
        run = cell.paragraphs[0].runs[0]
        run.bold = True
    for i, (col1, col2) in enumerate(rows_data):
        row = table.rows[i + 1].cells
        row[0].text = col1
        row[1].text = col2
    doc.add_paragraph()


def generate_governance_doc(domain, org_name, regulatory_drivers, output_path):
    tmpl_key = None
    for k in DOMAIN_TEMPLATES:
        if k.lower() in domain.lower() or domain.lower() in k.lower():
            tmpl_key = k
            break

    tmpl = DOMAIN_TEMPLATES.get(tmpl_key, GENERIC_TEMPLATE)

    def fmt(s):
        return s.format(org_name=org_name, domain=domain)

    doc = Document()

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"{domain} Policy, Standard & Guideline")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run(f"{org_name}  |  SCRMS-PDCA Governance Documentation  |  Version 1.0 DRAFT")
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_metadata_table(doc, org_name, domain, "1.0 DRAFT", regulatory_drivers)

    # ── SECTION 1: POLICY ─────────────────────────────────────────────────────
    add_section(doc, "SECTION 1: POLICY", level=1)
    doc.add_paragraph(
        "A policy expresses management intent — what must be achieved and why. "
        "It is enforced; violations have consequences. Policies should be written in business language, "
        "approved by executive leadership, and reviewed annually."
    ).runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    add_section(doc, "1.1 Policy Statement")
    doc.add_paragraph(fmt(tmpl["policy_statement"]))

    add_section(doc, "1.2 Scope")
    doc.add_paragraph(fmt(tmpl["scope"]))

    add_section(doc, "1.3 Policy Requirements")
    add_bulleted_list(doc, tmpl["policy_requirements"])

    add_section(doc, "1.4 Exceptions")
    doc.add_paragraph(
        f"Exceptions to this policy must be submitted in writing to the [CISO / Security Governance function]. "
        f"All exceptions require: (1) business justification, (2) compensating controls, "
        f"(3) formal risk acceptance by an authorized approver, and (4) a defined remediation timeline. "
        f"Exception requests are tracked in the {org_name} exception register."
    )

    add_section(doc, "1.5 Policy Enforcement")
    doc.add_paragraph(
        "Violations of this policy may result in disciplinary action, up to and including termination of employment "
        "or contract. Suspected violations should be reported to [CISO / Security team contact]."
    )

    doc.add_page_break()

    # ── SECTION 2: STANDARD ───────────────────────────────────────────────────
    add_section(doc, "SECTION 2: STANDARD", level=1)
    doc.add_paragraph(
        "A standard provides specific, measurable requirements that operationalize the policy. "
        "Standards tell people exactly what to do and are enforced. They are more technical and "
        "specific than policies, and should be reviewed when technology or practices change."
    ).runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    add_section(doc, "2.1 Standard Controls")
    add_two_col_table(
        doc,
        [(ctrl, fmt(desc)) for ctrl, desc in tmpl["standard_controls"]],
        "Control Area",
        "Requirement"
    )

    add_section(doc, "2.2 Compliance Mapping")
    doc.add_paragraph(
        f"The following table maps these standards to applicable compliance obligations "
        f"identified in Principle #1 (Establish Context):"
    )
    reg_list = [r.strip() for r in regulatory_drivers.split(",")] if regulatory_drivers else ["[Regulation 1]", "[Framework 2]"]
    add_two_col_table(
        doc,
        [(reg, f"[Identify specific clauses/controls from {reg} that this standard addresses]") for reg in reg_list],
        "Regulatory/Framework Driver",
        "Mapped Clause / Control Reference"
    )

    doc.add_page_break()

    # ── SECTION 3: GUIDELINES ─────────────────────────────────────────────────
    add_section(doc, "SECTION 3: GUIDELINES", level=1)
    doc.add_paragraph(
        "Guidelines are recommended (non-mandatory) practices. They provide flexibility where standards "
        "don't prescribe an exact method, or where best practices exceed the minimum required by standards. "
        "Teams should consider guidelines when designing implementations."
    ).runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    add_section(doc, "3.1 Recommended Practices")
    add_bulleted_list(doc, tmpl["guidelines"])

    doc.add_page_break()

    # ── SECTION 4: CONTROL OBJECTIVES ─────────────────────────────────────────
    add_section(doc, "SECTION 4: CONTROL OBJECTIVES", level=1)
    doc.add_paragraph(
        "Control objectives define the outcomes that must be achieved. They link this governance documentation "
        "to the organization's control set and define the scope of implementation (which PPTDF categories apply). "
        "Control objectives are the interface between governance and audit/assessment."
    ).runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    add_section(doc, "4.1 Control Objectives")
    add_bulleted_list(doc, tmpl["control_objectives"])

    add_section(doc, "4.2 PPTDF Applicability")
    add_two_col_table(
        doc,
        [
            ("People",      "[Does this control objective apply to people? What roles/groups?]"),
            ("Processes",   "[Does this control objective apply to business/IT processes? Which ones?]"),
            ("Technologies","[Does this control objective apply to specific systems, apps, or infrastructure?]"),
            ("Data",        "[Does this control objective apply to specific data types or classifications?]"),
            ("Facilities",  "[Does this control objective apply to physical locations or environments?]"),
        ],
        "PPTDF Category",
        "Applicability Notes"
    )

    doc.add_page_break()

    # ── SECTION 5: DOCUMENT CONTROL ───────────────────────────────────────────
    add_section(doc, "SECTION 5: DOCUMENT CONTROL", level=1)
    add_two_col_table(
        doc,
        [
            ("Document Owner",      "[CISO / Security Governance Role]"),
            ("Approved By",         "[Executive Approver]"),
            ("Review Frequency",    "Annual (or upon significant change to regulatory drivers or technology)"),
            ("Next Review Date",    f"{date.today().replace(year=date.today().year + 1).strftime('%B %Y')}"),
            ("Distribution",        "All personnel within scope"),
            ("Classification",      "Internal"),
        ],
        "Field",
        "Value"
    )

    add_section(doc, "5.1 Revision History")
    add_two_col_table(
        doc,
        [("1.0", f"{date.today().strftime('%B %d, %Y')} — Initial draft (SCRMS-PDCA Governance Generator)")],
        "Version",
        "Change Description"
    )

    doc.save(output_path)
    print(f"\n✅ Governance document saved to: {output_path}")
    print("\nDocument sections:")
    print("  Section 1: Policy (statement, scope, requirements, exceptions, enforcement)")
    print("  Section 2: Standard (controls + compliance mapping)")
    print("  Section 3: Guidelines (recommended practices)")
    print("  Section 4: Control Objectives (outcomes + PPTDF scope)")
    print("  Section 5: Document Control (owner, approvals, revision history)")


def main():
    parser = argparse.ArgumentParser(
        description="Generate a SCRMS governance document (Policy + Standard + Guideline)"
    )
    parser.add_argument("--domain", required=True,
                        help="Control domain (e.g. 'Access Management', 'Data Protection')")
    parser.add_argument("--org-name", default="[Organization Name]",
                        help="Organization name")
    parser.add_argument("--regulatory-drivers", default=None,
                        help="Comma-separated compliance drivers (e.g. 'HIPAA,SOC 2,ISO 27001')")
    parser.add_argument("--output", default=None,
                        help="Output .docx path (default: <domain>_governance.docx)")
    args = parser.parse_args()

    output = args.output or f"{args.domain.replace(' ', '_').lower()}_governance.docx"
    generate_governance_doc(args.domain, args.org_name, args.regulatory_drivers, output)


if __name__ == "__main__":
    main()
